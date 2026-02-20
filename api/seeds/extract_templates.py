#!/usr/bin/env python3
"""
Extract 41 form templates from the DOCX file and save as JSON.

The DOCX has 42 tables, but table index 7 is a continuation sub-table
belonging to MODELO 7 (table index 6). We merge it into MODELO 7's content.
"""

import json
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


DOCX_PATH = "/Users/lindapeixoto/My Apps/igj/sandbox/docs/modulos/Processo Administrativo - Formulários.docx"
OUTPUT_PATH = "/Users/lindapeixoto/My Apps/igj/sandbox/api/seeds/pecas_templates.json"

SUBTABLE_INDEX = 7
PARENT_TABLE_INDEX = 6


def parse_title(raw_title):
    raw_title = raw_title.strip()
    m = re.match(
        r'^(ANEXO|MODELO)\s+(\d+)\s*[\u2013\-\u2014]\s*(.+)$',
        raw_title, re.IGNORECASE
    )
    if m:
        prefix = m.group(1).upper()
        num = m.group(2)
        titulo_part = m.group(3).strip()
        numero = f"{prefix} {num}"
        return numero, titulo_part, titulo_part
    return raw_title, raw_title, raw_title


def cell_to_html(cell):
    """
    Convert a DOCX cell to HTML. We process at the run level,
    collecting segments that are either:
      - underscored text (placeholder blanks)
      - parenthesized italic/underline text (field labels)
      - normal formatted text
    
    We then assemble the HTML and do field-label merging in post-processing.
    """
    html_parts = []

    for para in cell.paragraphs:
        align = ""
        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            align = ' style="text-align:center"'
        elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            align = ' style="text-align:right"'
        elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            align = ' style="text-align:justify"'

        # Collect segments: each is (text, is_bold, is_italic, is_underline)
        segments = []
        for run in para.runs:
            text = run.text
            if not text:
                continue
            segments.append((text, bool(run.bold), bool(run.italic), bool(run.underline)))

        # Now build HTML from segments, detecting placeholder patterns
        # A placeholder field is: underscores + optional (label) + underscores
        # Labels appear as italic+underline runs within parentheses
        inline = _segments_to_html(segments)

        # If no runs but has text
        if not inline and para.text.strip():
            text = para.text
            text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            inline = text

        html_parts.append(f"<p{align}>{inline}</p>")

    return "\n".join(html_parts)


def _segments_to_html(segments):
    """
    Process run segments into HTML. We detect placeholder patterns:
    - Runs of 3+ underscores -> [CAMPO] marker
    - Parenthesized text (especially italic+underline) between underscore runs -> field label
    
    Strategy: first convert to interim HTML, then do regex cleanup.
    """
    parts = []
    for text, bold, italic, underline in segments:
        escaped = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        
        # Check if this segment is purely underscores (placeholder blank)
        if re.match(r'^_+$', text) and len(text) >= 2:
            parts.append('{{BLANK}}')
            continue
        
        # Check if this is a parenthesized label (italic, often underline)
        # These look like "(label text)" or "( label text )"
        # They appear between underscore runs
        stripped = text.strip()
        
        # Apply formatting
        result = escaped
        if bold:
            result = f"<strong>{result}</strong>"
        if italic:
            result = f"<em>{result}</em>"
        if underline:
            result = f"<u>{result}</u>"
        
        parts.append(result)
    
    raw_html = "".join(parts)
    
    # Now post-process to convert {{BLANK}} markers and labels
    # First: collapse consecutive {{BLANK}} markers
    raw_html = re.sub(r'(\{\{BLANK\}\}\s*)+', '{{BLANK}}', raw_html)
    
    # Pattern: {{BLANK}} followed by parenthesized label followed by {{BLANK}}
    # The label might have <u>, <em> tags around it
    # E.g.: {{BLANK}}<u> (</u><u><em>label text</em></u><u>)</u>{{BLANK}}
    # Or: {{BLANK}}<u>(</u><u><em>label</em></u><u>)</u>{{BLANK}}
    
    # Extract labeled fields: {{BLANK}} + (label) + {{BLANK}}
    def replace_labeled_field(m):
        inner = m.group(1)
        # Strip all HTML tags to get the plain label text
        label = re.sub(r'<[^>]+>', '', inner).strip()
        # Remove surrounding parentheses and whitespace
        label = label.strip('() \t')
        if label:
            return f'<span style="color:blue">[{label}]</span>'
        return '<span style="color:blue">[CAMPO]</span>'
    
    # Match {{BLANK}} + stuff with parens + {{BLANK}}
    # The "stuff with parens" is everything between the blanks that contains ( and )
    raw_html = re.sub(
        r'\{\{BLANK\}\}((?:<[^>]*>|\s)*\((?:(?:<[^>]*>|[^)])*)\)(?:<[^>]*>|\s)*)\{\{BLANK\}\}',
        replace_labeled_field,
        raw_html
    )
    
    # Any remaining {{BLANK}} markers become [CAMPO]
    raw_html = raw_html.replace('{{BLANK}}', '<span style="color:blue">[CAMPO]</span>')
    
    return raw_html


def main():
    doc = Document(DOCX_PATH)
    tables = doc.tables

    print(f"Total tables found: {len(tables)}")

    templates = []
    skip_indices = {SUBTABLE_INDEX}

    for i, table in enumerate(tables):
        if i in skip_indices:
            continue

        rows = table.rows
        if len(rows) < 2:
            print(f"  Skipping table {i} (only {len(rows)} row, no title)")
            continue

        raw_title = rows[0].cells[0].text.strip()
        numero, titulo, designacao = parse_title(raw_title)

        content_cell = rows[1].cells[0]
        html_content = cell_to_html(content_cell)

        if i == PARENT_TABLE_INDEX:
            subtable = tables[SUBTABLE_INDEX]
            sub_cell = subtable.rows[0].cells[0]
            sub_html = cell_to_html(sub_cell)
            html_content += "\n" + sub_html

        templates.append({
            "numero": numero,
            "titulo": titulo,
            "designacao": designacao,
            "conteudo_html": html_content
        })

    print(f"\nExtracted {len(templates)} templates")

    for t in templates:
        content_preview = t["conteudo_html"][:60].replace("\n", "\\n")
        print(f"  {t['numero']:15s} | {t['titulo'][:55]:55s} | {content_preview}...")

    with open(OUTPUT_PATH, "w", encoding="utf-8") as f:
        json.dump(templates, f, ensure_ascii=False, indent=2)

    print(f"\nSaved to {OUTPUT_PATH}")
    print(f"Total entries: {len(templates)}")


if __name__ == "__main__":
    main()
